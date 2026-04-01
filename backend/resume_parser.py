import pdfplumber
import re
from docx import Document
from io import BytesIO
import os

def extract_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf_bytes(file_bytes)
    elif ext == '.docx':
        return extract_text_from_docx_bytes(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")


def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^a-zA-Z0-9.,@+\- ]', '', text)
    return text.strip()
